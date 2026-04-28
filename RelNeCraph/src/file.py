import errors
import pandas as pd
from pathlib import Path
import sys
import os
from docx import Document
import pdfplumber
from bs4 import BeautifulSoup
from enum import Enum
from consts import seperator,supported_fmts,headers
import requests
class FMT(Enum):
    NULL=0
    CSV=1
    XLS=2
    PDF=3
    TXT=4
    HTML=5
    DOCS=6
    LOG=7
    SQL=8
    RAW=9
    COM=10


def read_specfile(path: str,fmt: FMT) -> tuple[str | None, None | Exception]:
    if fmt == FMT.COM:
        res = requests.get(path, timeout=5, headers=headers)
        if res.status_code == 200:
            html = res.text
        else:
            return ("",errors.RelNeException(str(res),E_HTTP))
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        ret = soup.get_text()
        return (ret,None) 

    ret = ""
    os_p = Path(path)
    if not os_p.exists():
        return (None,errors.RelNeException(f"Unable to parse {path} as a valid file of type {fmt}",errors.ERRNO.E_OS)) 
    
    if fmt == FMT.CSV:
        ret = pd.read_csv(path,delim=",")
        if not ret:
            return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT)) 
        buf = StringIO()
        return (ret,None,FMT.CSV)
    elif fmt == FMT.XLS: 
        data = pd.ExcelFile(path)
        if not ret:
            return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT)) 
        buf = StringIO()
        ret += "{"
        for sheet in data.sheet_names:
            df = xls.parse(sheet)
            buf = StringIO()
            df.to_csv(buf, index=False)
            ret += f"{{Sheet{sheet}{seperator}" + buf.getvalue() + f"}},"
        ret=ret[:-1]
        ret += "}"
        return (ret,None)
    elif fmt == FMT.PDF:
        with pdfplumber.open(path) as pdf:
            for i,page in enumerate(pdf.pages):
                ret += f"Page{i}: "+page.extract_text() + seperator
        return (ret,None)
    elif fmt == FMT.TXT:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                ret += line.strip()
        return (ret,None)
    elif fmt == FMT.HTML:
        with open(path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        ret = soup.get_text(separator="\n", strip=True)
        return (ret,None)
    elif fmt == FMT.DOCS:
        doc = Document(path)
        for x in doc.paragraphs:
            if x.text.strip():
                ret += f"Para:{x}{seperator}"
        if ret.endswith(seperator):
            ret = ret[:-len(seperator)]
        for table in doc.tables:
            ret += "{"
            for row in table.rows:       
                ret += "{" + f"{row}:"
                for cell in row.cells:
                    ret += f"{cell.text}{seperator}"
                if ret.endswith(seperator):
                    ret = ret[:-len(seperator)]
                ret += "},"
            if ret.endswith(","): ret[:-1]
            ret += "}" 
        return (ret,None)
       
    elif fmt == FMT.LOG:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                ret += line.strip() + "\n"
        return (ret,None) 

    elif fmt == FMT.SQL:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                if not line.startswith("//"):
                    ret += line.strip()+seperator 
        return (ret[:-6] if ret.endswith(seperator) else ret,None) 
    else:
        return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT)) 


def read_file(path: str) -> tuple[str | None, None | Exception, FMT]:
    if "http" in path:
        html = requests.get(path).text
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        ret = soup.get_text()
        return (ret,None,FMT.COM) 

    fmt = path.split(".")[-1]
    known = False
    ret = ""
    for i in supported_fmts:
        if not path.endswith(i):
            continue
        else:
            fmt = i 
            known = True
            break

    if not known:
        os_p = Path(path)
        if os_p.exists():
            return (path,None,FMT.RAW)
        else:
            return (None,errors.RelNeException(f"Unable to parse {path} as a valid file of type {fmt}",errors.ERRNO.E_OS),FMT.NULL) 
    
    
    if fmt == ".csv":
        ret = pd.read_csv(path,delim=",")
        if not ret:
            return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT),FMT.NULL) 
        buf = StringIO()
        return (ret,None,FMT.CSV)
    elif fmt == ".xls" or fmt == ".xlsx": 
        data = pd.ExcelFile(path)
        if not ret:
            return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT),FMT.NULL) 
        buf = StringIO()
        ret += "{"
        for sheet in data.sheet_names:
            df = xls.parse(sheet)
            buf = StringIO()
            df.to_csv(buf, index=False)
            ret += f"{{Sheet{sheet}{seperator}" + buf.getvalue() + f"}},"
        ret=ret[:-1]
        ret += "}"
        return (ret,None,FMT.XLS)
    elif fmt == ".pdf":
        with pdfplumber.open(path) as pdf:
            for i,page in enumerate(pdf.pages):
                ret += f"Page{i}: "+page.extract_text() + seperator
        return (ret,None,FMT.PDF)
    elif fmt == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                ret += line.strip()
        return (ret,None,FMT.TXT)
    elif fmt == ".html":
        with open(path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        ret = soup.get_text(separator="\n", strip=True)
        return (ret,None,FMT.HTML)

    elif fmt == ".docx" or fmt == ".doc":
        doc = Document(path)
        for x in doc.paragraphs:
            if x.text.strip():
                ret += f"Para:{x}{seperator}"
        if ret.endswith(seperator):
            ret = ret[:-len(seperator)]
        for table in doc.tables:
            ret += "{"
            for row in table.rows:       
                ret += "{" + f"{row}:"
                for cell in row.cells:
                    ret += f"{cell.text}{seperator}"
                if ret.endswith(seperator):
                    ret = ret[:-len(seperator)]
                ret += "},"
            if ret.endswith(","): ret[:-1]
            ret += "}" 
        return (ret,None,FMT.DOCS)
       
    elif fmt == ".log":
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                ret += line.strip() + "\n"
        return (ret,None,FMT.LOG) 

    elif fmt == ".sql":
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                if not line.startswith("//"):
                    ret += line.strip()+seperator 
        return (ret[:-6] if ret.endswith(seperator) else ret,None,FMT.SQL) 
    else:
        return (None,errors.RelNeException(f"Unable to parse as {fmt[1:]} file",errors.ERRNO.E_FMT)) 
